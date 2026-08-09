from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "ALE_1000_Task_Delivery_Plan_EDITABLE.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF2CC"
PALE_GREEN = "E2F0D9"
PALE_RED = "FCE4D6"
MID_GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "000000"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=11,
                 bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def set_style_font(style, latin="Calibri", east_asia="Microsoft YaHei", size=11,
                   color=BLACK, bold=None):
    style.font.name = latin
    style._element.get_or_add_rPr()
    rfonts = style._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def apply_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        set_row_cant_split(row)
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="B7B7B7", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_paragraph_border_bottom(paragraph, color=BLUE, size=12, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_body(doc, text="", bold_prefix=None, italic=False, keep_with_next=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE_YELLOW, accent="7F6000"):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(2)
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.2, bold=(first_col_bold and i == 0))
    apply_table_geometry(table, widths, indent_dxa=120)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(2)
    after.paragraph_format.space_after = Pt(0)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

doc.core_properties.title = "1,000 ALE-style Tasks Delivery Plan"
doc.core_properties.subject = "Editable execution plan for a client-commissioned ALE-style benchmark"
doc.core_properties.author = "Interview working draft"
doc.core_properties.comments = "Version 0.2 - explicit project baselines are marked for review"

# Standard business brief token map.
normal = doc.styles["Normal"]
set_style_font(normal, size=11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = doc.styles[name]
    set_style_font(style, size=size, color=color, bold=True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    st = doc.styles[list_name]
    set_style_font(st, size=11)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167
doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
doc.styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.75)
doc.styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.25)
doc.styles["List Number"].paragraph_format.left_indent = Inches(0.5)
doc.styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)

# Running header/footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("ALE-STYLE BENCHMARK DELIVERY PLAN  |  EDITABLE WORKING DRAFT")
set_run_font(hr, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Page ")
set_run_font(fr, size=9, color=MID_GRAY)
add_field(fp, "PAGE", "1")
fr2 = fp.add_run(" of ")
set_run_font(fr2, size=9, color=MID_GRAY)
add_field(fp, "NUMPAGES", "1")

# Memo masthead.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("DELIVERY DESIGN MEMO")
set_run_font(r, size=23, bold=True, color=BLACK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("1,000 ALE-style Tasks: Scope Definition and Executable Production Plan")
set_run_font(r, size=14, color="373737")

for label, value in (
    ("Purpose", "Interview take-home / client delivery design"),
    ("Version", "v0.2 - editable working draft"),
    ("Date", "2026-08-08"),
    ("Reference", "Agents' Last Exam, arXiv:2606.05405v2"),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(10)
set_paragraph_border_bottom(rule, color=BLUE, size=12, space=3)

add_callout(
    doc,
    "SCOPE LOCK / 可修改项 01",
    "本方案把客户的“1,000 条”定义为 1,000 个通过验收的可运行 instances。其内部结构固定为：960 个不同的专业 workflows 各产生 1 个主实例，另加 40 个关键流程的受控变体。该定义对齐 ALE v2 的 workflow 规模，但不声称复制 ALE 的 960 个原始 workflows。",
    fill=PALE_YELLOW,
    accent="7F6000",
)

add_heading(doc, "Executive decision", 1)
add_body(doc, "这不是一个“生产一千条题目”的标注项目，而是一项测量系统建设工程。最终验收对象是可从干净环境启动、由 Agent 独立完成、能产出明确 artifact，并可被校准评分器复验的任务实例。")
add_body(doc, "ALE v2 的统计口径是 960 个 workflows 与 1,490 个 instances，二者不可混用。由于 ALE 只公开约 150 个任务，客户项目无法也不应以“复刻 ALE 私有题库”为目标；可执行目标应是建立一套与 ALE 同量级的 960-workflow 自有覆盖图，并交付 1,000 个已验收实例。")

add_heading(doc, "1. Deliverable definition", 1)
add_table(
    doc,
    ["交付池", "数量", "计数单位", "用途与规则"],
    [
        ("开发与校准集", "100", "独立 workflows / instances", "供客户理解格式、联调 harness、校准评分器；不得用于最终排名。"),
        ("私有最终测试集", "760", "独立 workflows / instances", "冻结后仅由受控评测服务访问；作为正式比较与验收主集。"),
        ("滚动替换储备", "100", "独立 workflows / instances", "用于题目泄漏、软件升级或评分器失效后的版本替换。"),
        ("关键流程变体", "40", "instances", "从高价值、高风险流程中选 40 个，各增加 1 个输入或约束变体，用于鲁棒性检查。"),
        ("最终合计", "1,000", "accepted instances", "对应 960 个不同 workflows；全部通过质量门槛后方可计入。"),
        ("生产候选池", "1,250", "candidate instances", "按 80% 最终通过率设置 25% 生产缓冲；未通过项进入返工或淘汰，不混入交付数。"),
    ],
    [1900, 900, 1700, 4860],
    first_col_bold=True,
)

add_callout(
    doc,
    "RATIONALE",
    "“至少覆盖 960 workflows”在本方案中的含义是：建立 960 个彼此独立、具有不同工作目标或交付物的客户工作流，而不是给 960 个 ALE 私有题目换写 prompt。仅替换数字、公司名或输入文件，不产生新的 workflow，只能计为 instance variant。",
    fill=PALE_GREEN,
    accent="375623",
)

add_heading(doc, "2. Coverage allocation: 13 clusters / 55 subdomains", 1)
add_body(doc, "项目采用 ALE 的 13 个行业集群、55 个子领域作为一级覆盖框架，再由客户业务优先级进行二次加权。960 个 workflows 的分配在第 2 周冻结，方法如下：")
add_numbered(doc, "为每个子领域设置 10 个 workflow 的最低覆盖量，共 550 个，防止高流量领域挤压长尾能力。")
add_numbered(doc, "剩余 410 个 workflow 使用加权分配公式：客户重要性 40% + 经济价值 30% + 当前能力缺口 20% + 环境可实现性 10%。每项按 1-5 分评分。")
add_numbered(doc, "采用最大余数法把 410 个整数名额分配到 55 个子领域；任何单一子领域不得超过总量的 8%，除非客户书面批准。")
add_numbered(doc, "每个子领域至少包含 3 类任务：信息获取/分析、工具操作/产出、复核/决策；避免把覆盖等同于职业名称罗列。")

add_callout(
    doc,
    "可修改项 02 / CLIENT INPUT REQUIRED",
    "第 1 周客户需要提交业务优先级、禁用领域、软件许可边界和数据隐私等级。若客户未提供，项目仍按上述公式启动，但所有权重与分配结果会作为第一项正式签字件，而不是隐藏假设。",
    fill=PALE_YELLOW,
    accent="7F6000",
)

add_page_break(doc)

add_heading(doc, "3. Task contract: every instance is an executable package", 1)
add_body(doc, "每个实例必须包含下列 12 个版本化组件，缺一项即不得进入最终 QC：")
for item in (
    "workflow_id 与 instance_id：分别标识专业流程和可运行变体；",
    "domain / subdomain / role 标签及覆盖配额来源；",
    "专家署名记录、资历验证和冲突声明；",
    "面向 Agent 的任务说明、成功标准和禁止事项；",
    "输入文件包、数据来源、许可、隐私和脱敏记录；",
    "操作系统、软件版本、账号权限和工具 manifest；",
    "可复现的环境快照与启动/重置脚本；",
    "专家参考产物或可接受结果集合；",
    "evaluate() 评分器、rubric 与各项权重；",
    "正例、负例、边界例和 reward-hacking 测试；",
    "dry-run 轨迹、缺陷记录、修订历史和审批人；",
    "split 标签、版本号、访问权限、泄漏状态和退役规则。",
):
    add_bullet(doc, item)

add_heading(doc, "4. Seven quality gates", 1)
add_table(
    doc,
    ["Gate", "责任人", "强制产出", "通过标准"],
    [
        ("G0 需求与覆盖", "Program Lead + Domain Lead", "workflow brief、配额映射", "目标、交付物、用户价值和边界均明确；不与现有 workflow 重复。"),
        ("G1 专家与来源", "Domain Lead + Legal/Privacy", "专家记录、来源清单", "资历通过；数据授权、隐私和许可无阻断项。"),
        ("G2 任务工程", "Task Engineer", "环境、工具、输入、reference", "从干净快照连续 3 次可启动；工具可用率 100%。"),
        ("G3 Prompt-Verifier 对齐", "Evaluation Engineer", "覆盖矩阵、评分器、测试集", "每项要求均被评分；每项评分均有 prompt 或成功标准依据。"),
        ("G4 工程师 dry-run", "Independent QA", "完整轨迹、缺陷单", "非作者可独立完成；无缺失上下文、阻断工具或隐藏人工步骤。"),
        ("G5 对抗与校准", "Red-team QA + SME", "负例、边界例、作弊测试", "reference=1.0；空/损坏输出≤0.1；关键错误不能获得高分。"),
        ("G6 最终验收", "Acceptance Committee", "签字记录、split/version", "所有缺陷关闭；复跑通过；权限和私有集状态正确。"),
    ],
    [900, 1900, 2500, 4060],
    first_col_bold=True,
)

add_callout(
    doc,
    "COUNTING RULE",
    "只有通过 G6 的任务才计入 1,000 个 accepted instances。待 QC、返工中、环境不可复现或仅有 prompt 没有评分器的项目全部不计数。",
    fill=PALE_RED,
    accent="9C0006",
)

add_heading(doc, "5. Evaluation strategy by output type", 1)
add_table(
    doc,
    ["输出类型", "主评分机制", "人工介入", "验收要求"],
    [
        ("客观可验证", "确定性代码 / artifact checks", "专家定义规则；QA 抽样 10%", "同一输出重复评分一致；关键错误有明确扣分。"),
        ("多解但可 rubric 化", "结构化 rubric + 部分确定性检查", "2 名专家独立复核校准集；分歧时第 3 人裁决", "评分边界经至少 20 个样本校准；专家与自动评分趋势一致。"),
        ("高度主观", "专业人士盲法成对比较", "每个比较由 3 名合格评审；2/3 一致，否则裁决", "隐藏模型身份与生成顺序；报告胜率、置信区间及评审一致率。"),
        ("高风险专业判断", "自动检查 + 专家否决权", "领域专家审查全部关键失败", "严重安全、法律或事实错误触发 fail-closed，不被平均分掩盖。"),
    ],
    [1700, 2600, 2600, 2460],
    first_col_bold=True,
)

add_body(doc, "ALE 的专家主要参与上游任务定义、参考结果和 rubric QC，运行时以自动评分为主。因此，本方案不把专家盲评当成所有任务的第六个统一步骤；它只用于主观质量或高风险判断确实决定任务成功的类别。")

add_page_break(doc)

add_heading(doc, "6. 24-week production plan", 1)
add_table(
    doc,
    ["阶段", "时间", "累计验收", "关键动作与退出条件"],
    [
        ("Define", "W1-W2", "0", "冻结 13/55 覆盖图、960 workflow 配额、数据政策、任务 schema 与验收标准。"),
        ("Pilot", "W3-W4", "50", "每个行业集群至少 3 个试点；验证 Windows/Linux 环境、评分器与 QC 工时。"),
        ("Calibrate", "W5-W8", "200", "完成 150 个新增验收；基于缺陷数据调整模板、rubric 和专家培训。"),
        ("Scale I", "W9-W14", "530", "每周验收 55 个；覆盖所有 55 个子领域；开始冻结私有测试集。"),
        ("Scale II", "W15-W20", "860", "继续每周验收 55 个；补齐长尾配额，完成对抗测试和候选淘汰。"),
        ("Close", "W21-W22", "1,000", "完成最后 140 个验收；冻结 760 私有、100 储备、40 变体的版本与权限。"),
        ("Audit & Handoff", "W23-W24", "1,000", "独立抽样复跑、缺陷清零、文档/环境/评分器交接和客户验收。"),
    ],
    [1400, 1000, 1200, 5760],
    first_col_bold=True,
)

add_callout(
    doc,
    "可修改项 03 / SCHEDULE BASELINE",
    "24 周是本方案的承诺基线。若客户要求更短周期，必须增加并行专家与工程产能，或减少 960 个独立 workflows 的覆盖要求；不能通过压缩 G3-G6 的质量门槛换取进度。",
    fill=PALE_YELLOW,
    accent="7F6000",
)

add_heading(doc, "7. Team and throughput model", 1)
add_table(
    doc,
    ["角色", "基线配置", "职责 / 产能依据"],
    [
        ("Program & benchmark design", "1 Program Lead + 1 Benchmark Architect", "范围、覆盖、客户决策、版本和验收口径。"),
        ("Domain governance", "13 Domain Leads", "每个行业集群 1 名负责人；管理专家资历、工作流去重和专业质量。"),
        ("Expert authors", "42 名并行作者 + reserve pool", "稳定期每人每周提交约 1.5 个候选实例，支持约 62-65 个候选/周。"),
        ("Task engineering", "10 Task Engineers", "环境、软件、输入包、reference 与自动化启动。"),
        ("Evaluation engineering", "4 Evaluation Engineers", "evaluate()、rubric、测试集、评分器校准和 reward-hacking 检查。"),
        ("Independent QA", "6 QA Reviewers", "dry-run、对抗测试、缺陷分级、返工验收和批次抽样。"),
        ("Infrastructure", "3 Engineers", "Windows/Linux 镜像、运行编排、日志、权限、成本和可复现性。"),
        ("Legal / privacy / security", "2 shared reviewers", "数据许可、PII、客户政策、私有集访问与泄漏响应。"),
        ("Acceptance committee", "5 人跨职能委员会", "G6 最终签字；与任务作者分离。"),
    ],
    [2350, 2250, 4760],
    first_col_bold=True,
)

add_callout(
    doc,
    "可修改项 04 / STAFFING BASELINE",
    "人员配置由每周 62-65 个候选、约 80% 最终通过率和 50-55 个验收量反推。若专家提交率或一次通过率低于基线，先启用 reserve pool 和返工专班，再调整最终日期。",
    fill=PALE_YELLOW,
    accent="7F6000",
)

add_heading(doc, "8. Operating dashboard", 1)
add_table(
    doc,
    ["指标", "红线 / 目标", "管理动作"],
    [
        ("Accepted instances", "W4=50; W8=200; W14=530; W20=860; W22=1,000", "每周按 workflow 与 instance 双口径报数。"),
        ("Distinct workflow coverage", "≥960，且 55 个子领域均达最低配额", "重复或仅换输入的条目降级为 variant，不计新 workflow。"),
        ("最终通过率", "≥80%（1,250 candidates → 1,000 accepted）", "连续两周低于 75% 时暂停扩量，定位缺陷来源。"),
        ("干净环境启动成功率", "3/3 dry-runs；批量运行 ≥99%", "失败任务退出私有池并回到 G2。"),
        ("Prompt-verifier coverage", "100% 要求被覆盖；0 项无依据评分", "任何未覆盖项阻断 G3。"),
        ("严重 false accept", "0", "发现即冻结相关评分器与同模板任务，启动横向审计。"),
        ("私有集泄漏", "0", "撤下、换入 reserve、追踪访问日志并更新版本。"),
        ("主观评审一致率", "三人评审 2/3 一致率 ≥80%", "低于阈值时重写 rubric、复训评审并重新校准。"),
    ],
    [2600, 3300, 3460],
    first_col_bold=True,
)

add_page_break(doc)

add_heading(doc, "9. Governance and handoff", 1)
for item in (
    "公共/开发集、私有最终集、滚动储备和变体池分别存储，使用独立访问组；",
    "每次运行固定 agent harness、模型版本、系统提示、工具、环境、预算、重试策略和评分器版本；",
    "私有题目只通过评测服务下发，任务源码、隐藏 reference 与评分器不进入模型运行环境；",
    "所有任务和输入包使用内容哈希、版本标签与不可变发布清单；",
    "每月扫描泄漏线索、软件失效、许可变化和评分器异常；触发条件满足即从 reserve 替换；",
    "客户交接包含覆盖矩阵、任务包、环境镜像、评分器、QC 证据、访问清单、运行手册和变更日志。",
):
    add_bullet(doc, item)

add_heading(doc, "10. What the client receives", 1)
add_table(
    doc,
    ["交付物", "内容", "验收证据"],
    [
        ("1,000 accepted task packages", "960 个独立 workflows + 40 个变体", "G6 签字、版本清单、内容哈希。"),
        ("Coverage map", "13 clusters / 55 subdomains / 960 workflow quotas", "配额公式、客户权重、缺口与完成状态。"),
        ("Executable environments", "Windows/Linux 镜像、软件与权限 manifest", "3 次干净启动记录、批量运行成功率。"),
        ("Calibrated evaluators", "确定性评分器、rubric、盲评协议与测试集", "正/负/边界例、false-accept 审计。"),
        ("Audit trail", "来源、专家、dry-run、缺陷、返工与审批记录", "可追溯到 workflow_id / instance_id。"),
        ("Private-set governance", "访问、版本、泄漏、轮换与退役制度", "权限清单、轮换储备、响应演练。"),
    ],
    [2400, 3900, 3060],
    first_col_bold=True,
)

add_callout(
    doc,
    "FINAL POSITION",
    "客户购买的不是一千个 prompt，而是一个包含 1,000 个已验收实例、960 个专业工作流、可运行环境、校准评分器、质量证据和私有集治理的完整测量系统。",
    fill=LIGHT_BLUE,
    accent=DARK_BLUE,
)

add_heading(doc, "Review markers / 可直接提出修改的位置", 1)
add_table(
    doc,
    ["标记", "当前锁定值", "可修改内容"],
    [
        ("可修改项 01", "1,000 instances = 960 workflows + 40 variants", "工作流与变体的数量关系、最终计数单位。"),
        ("可修改项 02", "55 子领域最低 10 个 + 410 加权分配", "客户权重、禁用领域、软件与数据边界。"),
        ("可修改项 03", "24 周", "里程碑、并行批次和客户验收窗口。"),
        ("可修改项 04", "42 位活跃专家作者及配套工程/QA 团队", "人员规模、内外部比例与并行度。"),
        ("可修改项 05", "验收 KPI 与主观评审阈值", "通过率、抽样率、一致率与 fail-closed 条件。"),
    ],
    [1700, 3900, 3760],
    first_col_bold=True,
)

add_heading(doc, "Sources and evidence boundary", 1)
p = add_body(doc, "1. Agents' Last Exam, arXiv:2606.05405v2: ")
add_hyperlink(p, "Version-pinned paper", "https://arxiv.org/html/2606.05405v2")
p = add_body(doc, "2. Official ALE repository: ")
add_hyperlink(p, "Official GitHub implementation", "https://github.com/rdi-berkeley/agents-last-exam")
p = add_body(doc, "3. Nick Heiner, Surge AI, When Will The Benchmaxxing Plague End?: ")
add_hyperlink(p, "Conference talk video", "https://www.youtube.com/watch?v=-npY6XjM8CQ")
p = add_body(doc, "4. Surge AI, Hemingway-bench methodology: ")
add_hyperlink(p, "Writing benchmark methodology", "https://surgehq.ai/blog/hemingway-bench-ai-writing-leaderboard")

add_body(doc, "Evidence boundary: ALE v2 reports 960 workflows and 1,490 instances, including 150 public, 1,017 private and 323 pending QC. The 1,000-task delivery allocation, 24-week schedule, staffing plan, pool split and acceptance thresholds in this memo are proposed project decisions created for the client scenario; they are not claims made by the ALE authors.", italic=True)

doc.save(OUT)
print(OUT.name)
