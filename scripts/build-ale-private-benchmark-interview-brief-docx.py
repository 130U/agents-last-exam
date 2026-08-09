#!/usr/bin/env python3
"""Build the <=10-page interview brief from its Markdown source.

Design system: standard_business_brief with a named ``interview_brief_dense``
override for 0.78-inch margins, 9.6-point body copy, and compact tables. The
override is intentional because the brief has a hard ten-page reading budget.
The long-form Markdown remains the evidence layer.
"""

from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
BASE_BUILDER = ROOT / "build-ale-private-clone-delivery-report-draft-v3-docx.py"


def load_base_builder():
    spec = importlib.util.spec_from_file_location("ale_v3_docx_builder", BASE_BUILDER)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load base builder: {BASE_BUILDER}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


base = load_base_builder()

NAVY = "17365D"
BLUE = "245B8A"
CYAN = "168A8B"
DARK = "1F2937"
MID = "5B6673"
LIGHT = "F2F4F7"
PALE_CYAN = "E8F7F5"
WHITE = "FFFFFF"


def set_style_font(style, latin="Arial", east_asia="Microsoft YaHei") -> None:
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:ascii"), latin)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4.2)
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 15.0, NAVY, 12, 5.5),
        ("Heading 2", 11.8, BLUE, 8, 4),
        ("Heading 3", 10.5, NAVY, 6, 3),
    ):
        style = styles[name]
        set_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style)
        style.font.size = Pt(9.4)
        style.paragraph_format.space_after = Pt(2.2)
        style.paragraph_format.left_indent = Inches(0.24)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.line_spacing = 1.02


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)


def add_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("PRIVATE ALE-STYLE BENCHMARK  /  INTERVIEW BRIEF")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=7.2, color=MID)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("PAGE ")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=7.2, color=MID)
    base.add_page_field(fp)


def add_cover(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(6)
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("AGENTS’ LAST EXAM  /  DELIVERY RECOMMENDATION")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=8.6, color=CYAN)
    run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("客户私有 ALE-style Benchmark")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=24, color=NAVY)
    run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run("960–1,000 Workflows  /  1,490+ Runnable Instances")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=13.5, color=BLUE)
    run.bold = True

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.autofit = False
    callout.columns[0].width = Inches(6.94)
    cell = callout.cell(0, 0)
    base.set_cell_shading(cell, PALE_CYAN)
    base.set_cell_left_border(cell, color=CYAN, size="20")
    base.set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    base.add_inline(
        p,
        "**建议：**复刻 ALE 的测量架构，而不是公开题面；为单一客户交付 960 个 final-accepted workflows、1,490 个 final-QC instances，外部公开比例为零。",
        base_size=10.1,
        color=NAVY,
    )

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Inches(1.25)
    meta.columns[1].width = Inches(5.69)
    rows = (
        ("文档", "Interview brief v1；结论层，详细推导以 GitHub 长报告为准"),
        ("冻结日", "2026-08-09"),
        ("ALE 基线", "arXiv 2606.05405v2；Git 1e615e4…；HF a8c1fd1…"),
        ("Hook 来源", "2026-08-02 发布的官方录像；现场演讲日期为 2026-06-30"),
    )
    for idx, (label, value) in enumerate(rows):
        for c in meta.rows[idx].cells:
            base.set_cell_margins(c, top=55, start=0, bottom=55, end=90)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = meta.cell(idx, 0).paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        base.set_run_font(lr, name="Arial", east_asia="Microsoft YaHei", size=8.5, color=MID)
        lr.bold = True
        vp = meta.cell(idx, 1).paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        vr = vp.add_run(value)
        base.set_run_font(vr, name="Arial", east_asia="Microsoft YaHei", size=8.5, color=DARK)
    base.set_table_borders(meta, color=WHITE, size="0")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Operating principle")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=8.2, color=CYAN)
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Build a private measurement system—not an annotation batch.")
    base.set_run_font(run, name="Arial", east_asia="Microsoft YaHei", size=12.5, color=NAVY)
    run.bold = True

    doc.add_page_break()


def add_contents(_doc: Document, _headings: list[str]) -> None:
    """The short brief omits a standalone contents page to preserve page budget."""


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    base.set_table_borders(table, color="D7DEE7", size="4")

    width = Inches(6.94)
    if cols == 3 and rows[0][0] in {"顶层覆盖域", "Gate", "Manifest"}:
        if rows[0][0] == "顶层覆盖域":
            widths = [Inches(4.05), Inches(1.35), Inches(1.54)]
        elif rows[0][0] == "Gate":
            widths = [Inches(0.85), Inches(2.40), Inches(3.69)]
        else:
            widths = [Inches(1.35), Inches(2.05), Inches(3.54)]
    else:
        widths = [int(width / cols)] * cols
    for idx, col in enumerate(table.columns):
        col.width = widths[idx]

    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        base.set_repeat_table_rows_no_split(row)
        if r_idx == 0:
            base.repeat_table_header(row)
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            base.set_cell_margins(cell, top=55, start=85, bottom=55, end=85)
            if r_idx == 0:
                base.set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                base.set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            value = values[c_idx] if c_idx < len(values) else ""
            base.add_inline(p, value, base_size=7.7, color=WHITE if r_idx == 0 else DARK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)


def build(md_path: Path, out_path: Path) -> None:
    base.configure_styles = configure_styles
    base.configure_sections = configure_sections
    base.add_header_footer = add_header_footer
    base.add_cover = add_cover
    base.add_contents = add_contents
    base.add_table = add_table
    base.build_from_markdown(md_path, out_path)

    doc = Document(out_path)
    doc.core_properties.title = "客户私有 ALE-style Benchmark：交付建议书"
    doc.core_properties.subject = "960–1,000 workflows / 1,490+ runnable instances"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = "External interview brief; detailed evidence is linked in the repository."
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.input.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
