#!/usr/bin/env python3
"""Build the additive Draft v2 ALE workflow-package report as a polished DOCX.

The Markdown file remains the source of truth. This builder intentionally supports
the Markdown subset used by that report: headings, paragraphs, emphasis, links,
block quotes, lists, pipe tables, fenced code, and horizontal rules.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "245B8A"
CYAN = "19A7A8"
DARK = "1F2937"
MID = "52606D"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_CYAN = "E8F7F5"
WHITE = "FFFFFF"
RED = "9C2F2F"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_left_border(cell, color=CYAN, size="22") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_rows_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, label: str, target: str):
    part = paragraph.part
    relation_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([r_fonts, color, underline])
    new_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    new_run.append(text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, base_size=10.5, color=DARK, italic=False) -> None:
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=base_size, color=color)
            run.italic = italic
        token = match.group(0)
        if token.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=color)
            run.bold = True
            run.italic = italic
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=max(8.5, base_size - 0.5), color=NAVY)
            run.font.highlight_color = None
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=color)
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.left_indent = Inches(0.26)
        style.paragraph_format.first_line_indent = Inches(-0.18)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.34)
        section.footer_distance = Inches(0.34)


def add_header_footer(section) -> None:
    section.different_first_page_header_footer = False

    def populate_header(header):
        hp = header.paragraphs[0]
        hp.paragraph_format.space_after = Pt(0)

    def populate_footer(footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("PAGE ")
        set_run_font(run, size=7.5, color=MID)
        add_page_field(p)

    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def add_cover(doc: Document) -> None:
    mast = doc.add_paragraph()
    mast.paragraph_format.space_before = Pt(8)
    mast.paragraph_format.space_after = Pt(28)
    run = mast.add_run("UNIPAT  /  AGENTS’ LAST EXAM")
    set_run_font(run, size=9.5, color=CYAN)
    run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("从 Benchmaxxing 到\n可运行专业工作")
    set_run_font(run, size=26, color=NAVY)
    run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("1,000 个 ALE-style Workflow Packages\n生产与交付方案")
    set_run_font(run, size=15, color=BLUE)
    run.bold = True

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.autofit = False
    callout.columns[0].width = Inches(6.75)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_CYAN)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("DECISION DRAFT  ·  ADDITIVE-ONLY  ·  PILOT-CALIBRATED")
    set_run_font(run, size=9.5, color=NAVY)
    run.bold = True

    doc.add_paragraph()
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Inches(1.35)
    meta.columns[1].width = Inches(5.4)
    rows = (
        ("用途", "UniPat 面试作业 / 内部技术决策报告"),
        ("文档状态", "可交付初稿 v2；已吸收第一轮委托方反馈"),
        ("研究冻结日", "2026-08-09"),
        ("ALE 论文", "arXiv 2606.05405v2"),
        ("代码快照", "1e615e456de7cef57706680613cb80ee13c7fc76"),
        ("数据快照", "a8c1fd174a1f6cfa76526572a2e3ebece1276be2"),
    )
    for idx, (label, value) in enumerate(rows):
        for cell in meta.rows[idx].cells:
            set_cell_margins(cell, top=70, start=0, bottom=70, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p1 = meta.cell(idx, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9, color=MID)
        r1.bold = True
        p2 = meta.cell(idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=9, color=DARK)
    set_table_borders(meta, color=WHITE, size="0")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Operating principle")
    set_run_font(run, size=8.5, color=CYAN)
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Manage this as a private measurement-system build, not an annotation batch.")
    set_run_font(run, size=13, color=NAVY)
    run.bold = True

    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("内容导航")
    intro = doc.add_paragraph()
    add_inline(intro, "正文按决策顺序组织：先冻结测量对象和交付单位，再给出生产、验证、运行、统计、治理和扩产闭环。")
    for text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        set_run_font(run, size=10.2, color=BLUE)
        run.bold = True
    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.LEFT
    note.autofit = False
    note.columns[0].width = Inches(6.75)
    cell = note.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    add_inline(p, "阅读标签：**[F]** 来源事实；**[C]** 作者/机构主张；**[I]** 研究者推断；**[R]** 项目建议；**[P]** 待客户或 pilot 决定。", base_size=9.5)
    doc.add_page_break()


def add_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EC")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        content = line.strip().strip("|")
        rows.append([cell.strip() for cell in content.split("|")])
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    total = Inches(6.75)
    headers = rows[0]
    if cols == 2 and headers and headers[0].startswith("v2 顶层域"):
        widths = [Inches(1.80), Inches(4.95)]
    elif cols == 3 and headers and headers[0] == "Domain lane":
        widths = [Inches(1.25), Inches(2.45), Inches(3.05)]
    elif cols == 4 and headers and headers[0] == "Manifest":
        widths = [Inches(1.10), Inches(1.35), Inches(2.85), Inches(1.45)]
    else:
        widths = [int(total / cols)] * cols
    for idx, col in enumerate(table.columns):
        col.width = widths[idx]
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        set_repeat_table_rows_no_split(row)
        if r_idx == 0:
            repeat_table_header(row)
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            value = values[c_idx] if c_idx < len(values) else ""
            add_inline(p, value, base_size=8.2, color=WHITE if r_idx == 0 else DARK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.75)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code_lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.1, color=NAVY)
        if idx < len(code_lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_blockquote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.75)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_CYAN)
    set_cell_left_border(cell)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, base_size=10.5, color=NAVY, italic=True)


def add_heading(doc: Document, level: int, text: str, first_content_h1: list[bool]) -> None:
    if level == 1:
        first_content_h1[0] = True
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    if text.startswith("12.3 "):
        p.paragraph_format.page_break_before = True
    add_inline(p, text, base_size={1: 16, 2: 13, 3: 11.5}.get(level, 11), color=NAVY if level != 2 else BLUE)
    for run in p.runs:
        run.bold = True


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    add_inline(p, text)


def build_from_markdown(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    configure_sections(doc)
    for section in doc.sections:
        add_header_footer(section)

    add_cover(doc)
    contents = [line[2:].strip() for line in lines if line.startswith("# ") and not line.startswith("# 从 Benchmaxxing")]
    add_contents(doc, contents)

    # Cover absorbs the title, subtitle and metadata. Start at the reading guide.
    start = next(i for i, line in enumerate(lines) if line.strip() == "### 阅读标签")
    i = start
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    paragraph_lines: list[str] = []
    first_h1 = [False]

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            add_paragraph(doc, " ".join(s.strip() for s in paragraph_lines))
            paragraph_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, parse_table_rows(table_lines))
            table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            flush_table()
            in_code = True
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines.append(stripped)
            i += 1
            continue
        flush_table()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped == "---":
            flush_paragraph()
            add_rule(doc)
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            add_heading(doc, len(heading.group(1)), heading.group(2).strip(), first_h1)
            i += 1
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            add_blockquote(doc, stripped[2:])
            i += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1), base_size=10.2)
            i += 1
            continue
        numbered = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(2.5)
            run = p.add_run(numbered.group(1) + ". ")
            set_run_font(run, size=10.2, color=BLUE)
            run.bold = True
            add_inline(p, numbered.group(2), base_size=10.2)
            i += 1
            continue
        paragraph_lines.append(stripped)
        i += 1

    flush_paragraph()
    flush_table()
    if in_code:
        add_code_block(doc, code_lines)

    # Apply section geometry and update-fields-on-open flag after content creation.
    configure_sections(doc)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = "从 Benchmaxxing 到可运行专业工作"
    core.subject = "1,000 个 ALE-style Workflow Packages 的生产与交付方案"
    core.author = "Theodore Ouyang"
    core.comments = "Additive-only deliverable draft v2; copied from v1 and extended from repository Markdown."

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_from_markdown(args.input.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
